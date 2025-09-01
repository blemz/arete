"""
Conversation export service for the Arete philosophical tutoring system.

Provides functionality for exporting chat sessions to various formats
including PDF, Markdown, JSON, HTML, and plain text with customizable
formatting options and philosophical content optimization.
"""

import os
import json
import tempfile
from datetime import datetime
from typing import List, Dict, Any, Optional, Union
import time
import uuid
from pathlib import Path

# Import models
from ..models.chat_session import ChatSession, ChatMessage, MessageType
from ..models.export_models import (
    ExportFormat,
    ExportOptions,
    ExportResult,
    ExportTemplate,
    ExportJob
)

# Optional imports for different export formats
try:
    from fpdf import FPDF
    FPDF_AVAILABLE = True
except ImportError:
    FPDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False


class ConversationExportService:
    """Service for exporting philosophical conversations to various formats."""
    
    def __init__(self):
        """Initialize the export service."""
        self.temp_dir = tempfile.gettempdir()
        self.export_jobs: Dict[str, ExportJob] = {}
        self.templates: Dict[str, ExportTemplate] = {}
        
        # Initialize default templates
        self._initialize_default_templates()
    
    def export_session(
        self,
        session: ChatSession,
        format: ExportFormat,
        options: ExportOptions,
        template: Optional[ExportTemplate] = None
    ) -> ExportResult:
        """Export a conversation session to the specified format."""
        start_time = time.time()
        
        try:
            # Validate format availability
            if not self._is_format_available(format):
                return ExportResult(
                    success=False,
                    format=format,
                    error_message=f"Export format {format.value} is not available. Missing required dependencies."
                )
            
            # Generate filename
            filename = self._generate_filename(session, format, options)
            file_path = os.path.join(
                options.output_directory or self.temp_dir,
                filename
            )
            
            # Ensure directory exists
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # Export based on format
            if format == ExportFormat.JSON:
                self._export_to_json(session, file_path, options)
            elif format == ExportFormat.MARKDOWN:
                self._export_to_markdown(session, file_path, options, template)
            elif format == ExportFormat.HTML:
                self._export_to_html(session, file_path, options, template)
            elif format == ExportFormat.TXT:
                self._export_to_txt(session, file_path, options)
            elif format == ExportFormat.PDF:
                self._export_to_pdf(session, file_path, options)
            elif format == ExportFormat.DOCX:
                self._export_to_docx(session, file_path, options)
            elif format == ExportFormat.CSV:
                self._export_to_csv(session, file_path, options)
            else:
                raise ValueError(f"Unsupported export format: {format.value}")
            
            # Calculate export statistics
            export_time = time.time() - start_time
            file_size = os.path.getsize(file_path) if os.path.exists(file_path) else None
            
            # Create successful result
            result = ExportResult(
                success=True,
                file_path=file_path,
                format=format,
                file_size=file_size,
                export_time=export_time,
                message_count=len(session.messages),
                export_timestamp=datetime.now()
            )
            
            # Add compression if requested
            if options.compress_output:
                compressed_path = self._compress_file(file_path)
                if compressed_path:
                    original_size = file_size
                    compressed_size = os.path.getsize(compressed_path)
                    result.file_path = compressed_path
                    result.file_size = compressed_size
                    result.compression_ratio = compressed_size / original_size if original_size > 0 else None
                    # Remove original file
                    os.unlink(file_path)
            
            return result
            
        except Exception as e:
            return ExportResult(
                success=False,
                format=format,
                error_message=str(e),
                export_time=time.time() - start_time
            )
    
    def bulk_export_sessions(
        self,
        sessions: List[ChatSession],
        format: ExportFormat,
        options: ExportOptions
    ) -> List[ExportResult]:
        """Export multiple sessions to the specified format."""
        results = []
        
        for session in sessions:
            result = self.export_session(session, format, options)
            results.append(result)
        
        return results
    
    def create_export_job(
        self,
        session_id: str,
        user_id: str,
        format: ExportFormat,
        options: ExportOptions,
        priority: int = 5
    ) -> ExportJob:
        """Create an asynchronous export job."""
        job_id = f"export_{uuid.uuid4().hex[:8]}"
        
        job = ExportJob(
            job_id=job_id,
            session_id=session_id,
            user_id=user_id,
            format=format,
            options=options,
            priority=priority
        )
        
        self.export_jobs[job_id] = job
        return job
    
    def get_export_job(self, job_id: str) -> Optional[ExportJob]:
        """Get export job by ID."""
        return self.export_jobs.get(job_id)
    
    def create_template(
        self,
        template_id: str,
        name: str,
        description: str,
        format: ExportFormat,
        options: ExportOptions,
        created_by: Optional[str] = None
    ) -> ExportTemplate:
        """Create a new export template."""
        template = ExportTemplate(
            template_id=template_id,
            name=name,
            description=description,
            format=format,
            options=options,
            created_by=created_by
        )
        
        self.templates[template_id] = template
        return template
    
    def get_template(self, template_id: str) -> Optional[ExportTemplate]:
        """Get export template by ID."""
        return self.templates.get(template_id)
    
    def list_templates(self, format: Optional[ExportFormat] = None) -> List[ExportTemplate]:
        """List available export templates."""
        templates = list(self.templates.values())
        
        if format:
            templates = [t for t in templates if t.format == format]
        
        return templates
    
    def _is_format_available(self, format: ExportFormat) -> bool:
        """Check if export format is available."""
        if format == ExportFormat.PDF and not FPDF_AVAILABLE:
            return False
        if format == ExportFormat.DOCX and not DOCX_AVAILABLE:
            return False
        if format in [ExportFormat.JSON, ExportFormat.MARKDOWN, ExportFormat.HTML, ExportFormat.TXT, ExportFormat.CSV]:
            return True
        return True
    
    def _generate_filename(self, session: ChatSession, format: ExportFormat, options: ExportOptions) -> str:
        """Generate filename for exported file."""
        # Clean title for filename
        clean_title = "".join(c for c in session.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        clean_title = clean_title.replace(' ', '_')[:50]  # Limit length
        
        # Add timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Apply custom prefix if provided
        prefix = options.filename_prefix or "arete_conversation"
        
        filename = f"{prefix}_{clean_title}_{timestamp}.{format.value}"
        
        return filename
    
    def _export_to_json(self, session: ChatSession, file_path: str, options: ExportOptions) -> None:
        """Export conversation to JSON format."""
        export_data = {
            "export_info": {
                "format": "json",
                "version": "1.0",
                "exported_at": datetime.now().isoformat(),
                "options": options.to_dict() if options.include_metadata else {}
            },
            "session": session.to_dict()
        }
        
        # Filter sensitive information if requested
        if options.anonymize_user:
            export_data["session"]["user_id"] = "anonymized"
            for message in export_data["session"]["messages"]:
                if "user_id" in message:
                    message["user_id"] = "anonymized"
        
        # Remove metadata if not requested
        if not options.include_metadata:
            export_data["session"].pop("metadata", None)
            for message in export_data["session"]["messages"]:
                message.pop("metadata", None)
        
        # Remove timestamps if not requested
        if not options.include_timestamps:
            for message in export_data["session"]["messages"]:
                message.pop("timestamp", None)
        
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, indent=2, ensure_ascii=False)
    
    def _export_to_markdown(self, session: ChatSession, file_path: str, options: ExportOptions, template: Optional[ExportTemplate] = None) -> None:
        """Export conversation to Markdown format."""
        lines = []
        
        # Title
        lines.append(f"# {session.title}")
        lines.append("")
        
        # Metadata
        if options.include_metadata:
            lines.append("## Session Information")
            lines.append(f"- **Session ID**: {session.session_id}")
            lines.append(f"- **Created**: {session.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
            lines.append(f"- **Updated**: {session.updated_at.strftime('%Y-%m-%d %H:%M:%S')}")
            lines.append(f"- **Status**: {session.status.value}")
            lines.append(f"- **Messages**: {len(session.messages)}")
            lines.append("")
        
        # Context
        if options.include_context and session.context:
            lines.append("## Discussion Context")
            if session.context.student_level:
                lines.append(f"- **Academic Level**: {session.context.student_level}")
            if session.context.philosophical_period:
                lines.append(f"- **Philosophical Period**: {session.context.philosophical_period}")
            if session.context.current_topic:
                lines.append(f"- **Current Topic**: {session.context.current_topic}")
            if session.context.learning_objectives:
                lines.append(f"- **Learning Objectives**: {', '.join(session.context.learning_objectives)}")
            lines.append("")
        
        # Messages
        lines.append("## Conversation")
        lines.append("")
        
        if not session.messages:
            lines.append("*No messages in this conversation.*")
        else:
            for message in session.messages:
                # Skip system messages unless requested
                if message.message_type == MessageType.SYSTEM and not options.include_system_messages:
                    continue
                
                # Message header
                role = "User" if message.message_type == MessageType.USER else "Assistant"
                timestamp_str = ""
                if options.include_timestamps:
                    timestamp_str = f" - {message.timestamp.strftime('%H:%M:%S')}"
                
                lines.append(f"### {role}{timestamp_str}")
                lines.append("")
                
                # Message content
                content = message.content
                if options.format_philosophical_quotes:
                    # Simple quote formatting (could be enhanced)
                    if message.message_type == MessageType.ASSISTANT and any(word in content.lower() for word in ["aristotle", "plato", "kant", "socrates"]):
                        content = f"> {content}"
                
                lines.append(content)
                lines.append("")
                
                # Citations
                if options.include_citations and message.citations:
                    lines.append("**Sources:**")
                    for citation in message.citations:
                        if options.highlight_citations:
                            lines.append(f"- **{citation}**")
                        else:
                            lines.append(f"- {citation}")
                    lines.append("")
        
        # Custom footer
        if options.custom_footer:
            lines.append("---")
            lines.append(options.custom_footer)
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
    
    def _export_to_html(self, session: ChatSession, file_path: str, options: ExportOptions, template: Optional[ExportTemplate] = None) -> None:
        """Export conversation to HTML format."""
        html_parts = []
        
        # HTML structure
        html_parts.append("<!DOCTYPE html>")
        html_parts.append("<html lang=\"en\">")
        html_parts.append("<head>")
        html_parts.append("    <meta charset=\"UTF-8\">")
        html_parts.append("    <meta name=\"viewport\" content=\"width=device-width, initial-scale=1.0\">")
        html_parts.append(f"    <title>{session.title}</title>")
        
        # CSS styles
        css = self._get_html_css(options, template)
        html_parts.append(f"    <style>{css}</style>")
        
        html_parts.append("</head>")
        html_parts.append("<body>")
        
        # Header
        if options.custom_header:
            html_parts.append(f"    <header><h1>{options.custom_header}</h1></header>")
        
        # Main content
        html_parts.append("    <main>")
        html_parts.append(f"        <h1>{session.title}</h1>")
        
        # Metadata
        if options.include_metadata:
            html_parts.append("        <div class=\"metadata\">")
            html_parts.append("            <h2>Session Information</h2>")
            html_parts.append("            <ul>")
            html_parts.append(f"                <li><strong>Session ID:</strong> {session.session_id}</li>")
            html_parts.append(f"                <li><strong>Created:</strong> {session.created_at.strftime('%Y-%m-%d %H:%M:%S')}</li>")
            html_parts.append(f"                <li><strong>Messages:</strong> {len(session.messages)}</li>")
            html_parts.append("            </ul>")
            html_parts.append("        </div>")
        
        # Messages
        html_parts.append("        <div class=\"conversation\">")
        
        if not session.messages:
            html_parts.append("            <p><em>No messages in this conversation.</em></p>")
        else:
            for message in session.messages:
                if message.message_type == MessageType.SYSTEM and not options.include_system_messages:
                    continue
                
                role_class = message.message_type.value
                timestamp_str = ""
                if options.include_timestamps:
                    timestamp_str = f"<span class=\"timestamp\">{message.timestamp.strftime('%H:%M:%S')}</span>"
                
                html_parts.append(f"            <div class=\"message {role_class}\">")
                html_parts.append(f"                <div class=\"message-header\">{message.message_type.value.title()}{timestamp_str}</div>")
                html_parts.append(f"                <div class=\"message-content\">{message.content}</div>")
                
                # Citations
                if options.include_citations and message.citations:
                    html_parts.append("                <div class=\"citations\">")
                    html_parts.append("                    <strong>Sources:</strong>")
                    html_parts.append("                    <ul>")
                    for citation in message.citations:
                        citation_class = "highlighted" if options.highlight_citations else ""
                        html_parts.append(f"                        <li class=\"{citation_class}\">{citation}</li>")
                    html_parts.append("                    </ul>")
                    html_parts.append("                </div>")
                
                html_parts.append("            </div>")
        
        html_parts.append("        </div>")
        html_parts.append("    </main>")
        
        # Footer
        if options.custom_footer:
            html_parts.append(f"    <footer>{options.custom_footer}</footer>")
        
        html_parts.append("</body>")
        html_parts.append("</html>")
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(html_parts))
    
    def _export_to_txt(self, session: ChatSession, file_path: str, options: ExportOptions) -> None:
        """Export conversation to plain text format."""
        lines = []
        
        # Title
        lines.append(session.title)
        lines.append("=" * len(session.title))
        lines.append("")
        
        # Metadata
        if options.include_metadata:
            lines.append("Session Information:")
            lines.append(f"  Session ID: {session.session_id}")
            lines.append(f"  Created: {session.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
            lines.append(f"  Messages: {len(session.messages)}")
            lines.append("")
        
        # Messages
        lines.append("Conversation:")
        lines.append("-" * 50)
        lines.append("")
        
        if not session.messages:
            lines.append("No messages in this conversation.")
        else:
            for message in session.messages:
                if message.message_type == MessageType.SYSTEM and not options.include_system_messages:
                    continue
                
                # Message header
                role = "User" if message.message_type == MessageType.USER else "Assistant"
                timestamp_str = ""
                if options.include_timestamps:
                    timestamp_str = f" [{message.timestamp.strftime('%Y-%m-%d %H:%M:%S')}]"
                
                lines.append(f"{role}:{timestamp_str}")
                lines.append(message.content)
                
                # Citations (if included)
                if options.include_citations and message.citations:
                    lines.append("")
                    lines.append("Sources:")
                    for citation in message.citations:
                        lines.append(f"  - {citation}")
                
                lines.append("")
                lines.append("-" * 30)
                lines.append("")
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
    
    def _export_to_pdf(self, session: ChatSession, file_path: str, options: ExportOptions) -> None:
        """Export conversation to PDF format."""
        if not FPDF_AVAILABLE:
            raise ImportError("FPDF library not available for PDF export")
        
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font('Arial', 'B', 16)
        
        # Title
        pdf.cell(0, 10, session.title, ln=True, align='C')
        pdf.ln(10)
        
        pdf.set_font('Arial', '', options.font_size)
        
        # Messages
        for message in session.messages:
            if message.message_type == MessageType.SYSTEM and not options.include_system_messages:
                continue
            
            # Message header
            role = "User" if message.message_type == MessageType.USER else "Assistant"
            timestamp_str = ""
            if options.include_timestamps:
                timestamp_str = f" - {message.timestamp.strftime('%H:%M:%S')}"
            
            pdf.set_font('Arial', 'B', options.font_size)
            pdf.cell(0, 8, f"{role}{timestamp_str}", ln=True)
            
            # Message content
            pdf.set_font('Arial', '', options.font_size)
            # Simple text wrapping (FPDF has limited text handling)
            words = message.content.split()
            line = ""
            for word in words:
                if pdf.get_string_width(line + word + " ") < 180:  # Page width minus margins
                    line += word + " "
                else:
                    if line:
                        pdf.cell(0, 6, line.strip(), ln=True)
                        line = word + " "
            if line:
                pdf.cell(0, 6, line.strip(), ln=True)
            
            pdf.ln(5)
        
        pdf.output(file_path)
    
    def _export_to_docx(self, session: ChatSession, file_path: str, options: ExportOptions) -> None:
        """Export conversation to DOCX format."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library not available for DOCX export")
        
        doc = Document()
        
        # Title
        title = doc.add_heading(session.title, 0)
        
        # Messages
        for message in session.messages:
            if message.message_type == MessageType.SYSTEM and not options.include_system_messages:
                continue
            
            # Message header
            role = "User" if message.message_type == MessageType.USER else "Assistant"
            timestamp_str = ""
            if options.include_timestamps:
                timestamp_str = f" - {message.timestamp.strftime('%H:%M:%S')}"
            
            header = doc.add_heading(f"{role}{timestamp_str}", level=2)
            
            # Message content
            doc.add_paragraph(message.content)
            
            # Citations
            if options.include_citations and message.citations:
                citations_para = doc.add_paragraph("Sources:")
                for citation in message.citations:
                    doc.add_paragraph(f"• {citation}", style='List Bullet')
        
        doc.save(file_path)
    
    def _export_to_csv(self, session: ChatSession, file_path: str, options: ExportOptions) -> None:
        """Export conversation to CSV format."""
        import csv
        
        with open(file_path, 'w', newline='', encoding='utf-8') as f:
            fieldnames = ['message_id', 'role', 'content', 'timestamp']
            if options.include_citations:
                fieldnames.append('citations')
            if options.include_metadata:
                fieldnames.extend(['provider', 'response_time', 'token_count'])
            
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            
            for message in session.messages:
                if message.message_type == MessageType.SYSTEM and not options.include_system_messages:
                    continue
                
                row = {
                    'message_id': message.message_id,
                    'role': message.message_type.value,
                    'content': message.content
                }
                
                if options.include_timestamps:
                    row['timestamp'] = message.timestamp.isoformat()
                
                if options.include_citations:
                    row['citations'] = '; '.join(message.citations)
                
                if options.include_metadata and message.metadata:
                    row.update({
                        'provider': message.metadata.get('provider', ''),
                        'response_time': message.metadata.get('response_time', ''),
                        'token_count': message.metadata.get('token_count', '')
                    })
                
                writer.writerow(row)
    
    def _get_html_css(self, options: ExportOptions, template: Optional[ExportTemplate] = None) -> str:
        """Get CSS styles for HTML export."""
        css = """
            body { font-family: Arial, sans-serif; margin: 20px; line-height: 1.6; }
            h1 { color: #2c3e50; border-bottom: 3px solid #3498db; }
            h2 { color: #34495e; }
            .message { margin: 20px 0; padding: 15px; border-radius: 8px; }
            .message.user { background-color: #e8f4fd; border-left: 4px solid #3498db; }
            .message.assistant { background-color: #f8f9fa; border-left: 4px solid #28a745; }
            .message-header { font-weight: bold; margin-bottom: 8px; }
            .timestamp { float: right; color: #6c757d; font-size: 0.9em; }
            .citations { margin-top: 10px; padding: 8px; background-color: #f1f1f1; border-radius: 4px; }
            .citations.highlighted { background-color: #fff3cd; }
            .metadata { background-color: #f8f9fa; padding: 15px; border-radius: 5px; margin: 20px 0; }
        """
        
        if template and template.css_styles:
            css += template.css_styles
        
        return css
    
    def _compress_file(self, file_path: str) -> Optional[str]:
        """Compress exported file using gzip."""
        try:
            import gzip
            compressed_path = file_path + '.gz'
            
            with open(file_path, 'rb') as f_in:
                with gzip.open(compressed_path, 'wb') as f_out:
                    f_out.writelines(f_in)
            
            return compressed_path
        except Exception:
            return None
    
    def _initialize_default_templates(self) -> None:
        """Initialize default export templates."""
        # Default Markdown template
        markdown_template = ExportTemplate(
            template_id="default_markdown",
            name="Default Markdown",
            description="Standard Markdown export with philosophical formatting",
            format=ExportFormat.MARKDOWN,
            options=ExportOptions(
                include_metadata=True,
                include_citations=True,
                include_timestamps=True,
                format_philosophical_quotes=True
            ),
            is_system_template=True,
            is_default=True
        )
        self.templates[markdown_template.template_id] = markdown_template
        
        # Default HTML template
        html_template = ExportTemplate(
            template_id="default_html",
            name="Default HTML",
            description="Standard HTML export with philosophical styling",
            format=ExportFormat.HTML,
            options=ExportOptions(
                include_metadata=True,
                include_citations=True,
                include_timestamps=True,
                highlight_citations=True
            ),
            is_system_template=True,
            is_default=True
        )
        self.templates[html_template.template_id] = html_template